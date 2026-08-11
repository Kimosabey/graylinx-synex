#!/usr/bin/env python3
"""Split the v4 reference document into one markdown file per chapter (task T2).

    python scripts/split_source.py            # write the chapters
    python scripts/split_source.py --check    # convert in memory and report fidelity only

`docs/00-source/` is read-only input (CLAUDE.md §3), so the .docx stays exactly where
it is. HANDOFF.md's version of T2 said to move it into `docs/90-archive/`; that
contradicts the read-only rule, and the rule wins — an archive is for superseded
editions, not for the source everything else derives from.

Fidelity is the whole point of this script, so it measures itself: word counts and
table counts are compared before and after, and a mismatch is an error rather than a
warning. `python-docx` yields paragraphs and tables as separate collections, which
loses their order — a table would land at the end of its chapter instead of inside the
argument it belongs to. So the body XML is walked directly instead.
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / "docs" / "00-source" / "Graylinx_Enterprise_AI_Platform_v4.docx"
PRODUCT = ROOT / "docs" / "10-product"
ARCHITECTURE = ROOT / "docs" / "20-architecture"

# Chapters 1-25 are PART ONE, 26-44 are PART TWO. The boundary is the document's own,
# and it is also CLAUDE.md hard rule 3: product first, architecture second.
ARCH_FROM = 26

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


def body_items(doc: Document):
    """Paragraphs and tables in document order, which the two collections lose."""
    for child in doc.element.body.iterchildren():
        if child.tag == f"{W}p":
            yield Paragraph(child, doc)
        elif child.tag == f"{W}tbl":
            yield Table(child, doc)


def inline(par: Paragraph) -> str:
    """Runs to markdown, keeping bold and italic and nothing else."""
    out = []
    for run in par.runs:
        text = run.text
        if not text:
            continue
        # Trailing/leading spaces must sit outside the markers or the emphasis breaks.
        lead = text[: len(text) - len(text.lstrip())]
        trail = text[len(text.rstrip()):]
        core = text.strip()
        if core:
            if run.bold and run.italic:
                core = f"***{core}***"
            elif run.bold:
                core = f"**{core}**"
            elif run.italic:
                core = f"*{core}*"
        out.append(f"{lead}{core}{trail}")
    return "".join(out)


def cell_text(cell) -> str:
    """A table cell, flattened. A pipe inside a cell would break the row."""
    parts = [inline(p) for p in cell.paragraphs]
    text = " ".join(p.strip() for p in parts if p.strip())
    return normalise(text.replace("|", "\\|").replace("\n", " "))


def render_table(tbl: Table) -> list[str]:
    rows = [[cell_text(c) for c in r.cells] for r in tbl.rows]
    if not rows:
        return []
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    head, *body = rows
    if not any(h.strip() for h in head):
        head = [f"col {i + 1}" for i in range(width)]
        body = rows
    out = ["", "| " + " | ".join(head) + " |",
           "|" + "|".join(["---"] * width) + "|"]
    out += ["| " + " | ".join(r) + " |" for r in body]
    out.append("")
    return out


# The source labels two of its own tables with tokens that collide with feature IDs:
# design levels L0-L6 against Alerts L1-L6, and release gates G0-G5 against the Control
# Plane G1-G6. The collision is worse than it looks, because L1-L6 and G1-G5 all *exist*
# in the register — so a reader or a script reading "L1" gets a confident wrong answer
# rather than an error. CLAUDE.md rule 7 forbids introducing such a prefix; this one
# arrived in the source, and docs/00-source/ is read-only, so it is normalised on the way
# out instead. Spelled out rather than re-prefixed, so no new prefix can collide later.
# The gate names are enumerated rather than pattern-matched on capitalisation: a loose
# rule would also rewrite a genuine reference to the Control Plane feature G6, turning a
# correct ID into a wrong one. Six labels, and only these six.
GATE_NAMES = "Design|Data|AI|Action|Production|Post-release"
# Two different L-series exist and must not be conflated with each other either:
# chapter 26 numbers *design levels* L0-L6, chapter 28 numbers *architecture layers*
# L0-L9. The first becomes "Level", the second "Layer", which is what each table's own
# header column already calls them.
NORMALISE = [
    (re.compile(r"\bL(\d)\s*(?=[—–-])"), r"Level \1 "),
    # normalise() runs after inline() has added emphasis, so the cell reads "**L0**"
    (re.compile(r"^(\**)L(\d)(\**)$"), r"\1Layer \2\3"),
    (re.compile(rf"\bG(\d)\s+(?={GATE_NAMES})"), r"Gate \1 — "),
]


def normalise(text: str) -> str:
    for pattern, repl in NORMALISE:
        text = pattern.sub(repl, text)
    return text


def slug(title: str) -> str:
    t = title.lower()
    t = re.sub(r"^\d+\.\s*", "", t)                 # the number lives in the prefix
    t = t.replace("&", " and ")
    t = re.sub(r"[\u2014\u2013]", "-", t)
    t = re.sub(r"[^a-z0-9]+", "-", t)
    return re.sub(r"-+", "-", t).strip("-")[:60]


def chapter_number(title: str) -> int | None:
    m = re.match(r"(\d+)\.\s", title)
    return int(m.group(1)) if m else None


def convert(doc: Document) -> list[dict]:
    """Walk the document once, cutting a new chapter at every Heading 1."""
    chapters: list[dict] = []
    current: dict | None = None
    skipping = False

    for item in body_items(doc):
        if isinstance(item, Table):
            if current and not skipping:
                current["lines"] += render_table(item)
                current["tables"] += 1
            continue

        style = item.style.name
        text = normalise(inline(item).strip())

        if style == "Heading 1":
            raw = item.text.strip()
            # "Contents" is a table of contents; a split makes it wrong by definition,
            # and each directory gets a generated index instead. PART ONE / PART TWO /
            # APPENDIX are dividers with no body of their own.
            if raw.lower() == "contents":
                skipping = True
                current = None
                continue
            skipping = False
            if raw in {"PART ONE", "PART TWO", "APPENDIX"}:
                current = None
                continue
            current = {"title": raw, "num": chapter_number(raw),
                       "lines": [f"# {text or raw}", ""], "tables": 0}
            chapters.append(current)
            continue

        if skipping or current is None:
            continue
        if not text:
            continue

        if style == "Heading 2":
            current["lines"] += ["", f"## {text}", ""]
        elif style == "Heading 3":
            current["lines"] += ["", f"### {text}", ""]
        elif style.startswith("List Bullet"):
            current["lines"].append(f"- {text}")
        elif style.startswith("List Number"):
            current["lines"].append(f"1. {text}")
        else:
            current["lines"] += [text, ""]

    return chapters


def tidy(lines: list[str]) -> str:
    out: list[str] = []
    for line in lines:
        if line == "" and out and out[-1] == "":
            continue
        out.append(line)
    return "\n".join(out).strip() + "\n"


def main() -> int:
    ap = argparse.ArgumentParser(description="Split the v4 reference into chapters")
    ap.add_argument("--check", action="store_true",
                    help="convert in memory and report fidelity only")
    args = ap.parse_args()

    if not SOURCE.exists():
        sys.exit(f"source not found: {SOURCE}")

    doc = Document(str(SOURCE))
    src_words = (sum(len(p.text.split()) for p in doc.paragraphs)
                 + sum(len(c.text.split()) for t in doc.tables
                       for r in t.rows for c in r.cells))
    src_tables = len(doc.tables)

    chapters = convert(doc)

    # Front matter has no chapter number, and must not borrow one: numbering it 00 and
    # 01 would put "Executive Summary" ahead of chapter 1 in a directory listing and
    # make it look like chapter 1. Everything unnumbered sits at 00 and sorts by title.
    plan = []
    for ch in chapters:
        n = ch["num"]
        if n is None:
            prefix, target = "00", PRODUCT
        else:
            prefix = f"{n:02d}"
            target = ARCHITECTURE if n >= ARCH_FROM else PRODUCT
        if ch["title"].lower().startswith("appendix"):
            prefix, target = "99", PRODUCT
        plan.append((target / f"{prefix}-{slug(ch['title'])}.md", ch))

    out_words = sum(len(tidy(ch["lines"]).split()) for _, ch in plan)
    out_tables = sum(ch["tables"] for _, ch in plan)

    print(f"source : {src_words:,} words, {src_tables} tables")
    print(f"output : {out_words:,} words, {out_tables} tables, "
          f"{len(plan)} files")

    problems = []
    if out_tables != src_tables:
        problems.append(f"table count {out_tables} != {src_tables}")
    # Markdown adds pipes and hyphens, so the word count rises. It must never fall.
    if out_words < src_words:
        problems.append(f"word count fell: {out_words:,} < {src_words:,}")

    dupes = [p for p, _ in plan if [q for q, _ in plan].count(p) > 1]
    if dupes:
        problems.append(f"duplicate filenames: {sorted({d.name for d in dupes})}")

    if problems:
        print("\nFAILED — " + "; ".join(problems))
        return 1

    if args.check:
        print("\nfidelity ok (nothing written)")
        return 0

    for target in (PRODUCT, ARCHITECTURE):
        target.mkdir(parents=True, exist_ok=True)
    for path, ch in plan:
        path.write_text(tidy(ch["lines"]), encoding="utf-8", newline="\n")

    for target, label in ((PRODUCT, "Product"), (ARCHITECTURE, "Architecture")):
        files = sorted(p for p in target.glob("*.md") if p.name != "README.md")
        index = [f"# {label} chapters", "",
                 f"Split from `docs/00-source/Graylinx_Enterprise_AI_Platform_v4.docx`",
                 "by `scripts/split_source.py`. The .docx remains the source of record;",
                 "these are the working copies. Re-run the script after it changes.", ""]
        for f in files:
            title = f.read_text(encoding="utf-8").splitlines()[0].lstrip("# ").strip()
            index.append(f"- [{title}]({f.name})")
        index.append("")
        (target / "README.md").write_text("\n".join(index), encoding="utf-8", newline="\n")

    print(f"\nwrote {len(plan)} chapters + 2 indexes")
    print(f"  {PRODUCT.relative_to(ROOT)}      "
          f"{len(list(PRODUCT.glob('*.md'))) - 1} chapters")
    print(f"  {ARCHITECTURE.relative_to(ROOT)} "
          f"{len(list(ARCHITECTURE.glob('*.md'))) - 1} chapters")
    return 0


if __name__ == "__main__":
    sys.exit(main())
