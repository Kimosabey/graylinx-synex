"""What the corpus **is** — the discovery half of `K1`/`K5` ingest, and the refusals.

`app/jobs/index_library.py` transcribes the checklist library into passages. This module
answers the question one step earlier and the harder one: *which documents in this repository
are genuinely retrievable knowledge?* Getting that wrong is not a missing feature, it is
`CLAUDE.md` §2.8 broken — **one source of truth per fact** — and the damage is silent, because
a corpus holding the same instruction twice returns confident, well-cited answers from both
copies and nothing in retrieval can tell they disagree.

So every candidate is decided **in this file, by name, with the reason** — including the ones
that are refused. A directory quietly not walked is indistinguishable from one somebody forgot,
which is the same argument `index_library.withheld_content` makes for the nine holding actions.

## What is ingested

| Source | Kind | Why it is knowledge |
|---|---|---|
| `docs/10-product/*.md` | `chapter` | The written product chapters — what Synex promises
  a user, in the form a person would cite |
| `docs/20-architecture/*.md` | `chapter` | The written architecture chapters. Same |
| `docs/00-source/WC_CHiller_FDD.docx` | `manual` | The water-cooled chiller FDD
  specification. The one source document with **no** derived form in this repository |
| the checklist library | `checklist` | 124 curated items, a 7-item fallback and 4
  differentials — transcribed in `app/domain/library/` and assembled by
  `index_library.library_documents()` |

## What is refused, and why each one would be a second source of truth

**The 78-page reference `.docx`.** `scripts/split_source.py` splits it into the 47 chapter
files above. Ingesting both puts every sentence in the corpus twice, and the `.docx` copy is
strictly worse: it has no per-chapter address, so its citation would name a 78-page document
and no place inside it — the `locator`-less citation `app/retrieval/chunking.py` exists to
prevent. The markdown is not a copy of the source, it *is* the source in the form that can be
cited.

**The Feature Review Pack, in both its formats.** Its checklist content is transcribed into
`app/domain/library/` and is already the `checklist` corpus, so ingesting the pack would make
three sources for one fact. The `.pdf` and the `.docx` are also the same document twice over.

**`docs/90-archive/`.** Superseded editions. Retrieving one is the precise failure `version`
exists to prevent: an instruction that was correct in an edition nobody follows any more.

**`mvp/`, `decisions/`, `brand/`, `CONTEXT.md`, `HANDOFF.md`, `CLAUDE.md`.** Documents about
*building* the product. A question an operator asks is not answered from the open-questions
register, and a passage from `HANDOFF.md` retrieved into an answer would read as a statement
about the plant.

**`app/domain/library/holding_actions.py`.** Withheld by `index_library.withheld_content()`
for a reason this module does not restate — it sits behind two gates and `DocumentChunk` has
one.

## The version question, which has two honest answers rather than one

`K5` promises the document **and** its version, and inventing one would make every citation in
the corpus name a revision that does not exist (`Q96`). But *no version anywhere* is not true
either, and asserting it would lose real provenance:

* the chapters split out of the reference document carry **edition 4**, which `CONTEXT.md` §8
  states and the source filename repeats. That is asserted by a source, not chosen here;
* the four architecture chapters written directly — the data model, the stack, the deployment
  shape and what was taken from Thermynx — belong to no edition and carry **nothing**;
* the FDD specification carries no revision number, no date and no edition. Also nothing.

Where there is nothing, `DocumentChunk.cite` prints `version not stated` rather than omitting
the field, so the absence is a sentence a reader can act on instead of a gap they fill in
themselves. Both states are present in the real corpus, deliberately, so neither is theoretical.

**Nothing here calls a model, and nothing here writes text.** A document is read off disk
verbatim; the only thing composed is the title, from the document's own H1 and the collection
it belongs to.
"""
from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from app.db.knowledge import digest_of

#: The kind a written chapter is tagged as — **never `sop`**, for the reason
#: `app/jobs/index_library.py` gives about the checklist library: `S4` restricts safety
#: retrieval to `sop`, so a kind that reached it would let a safety question be answered out of
#: a product chapter. A chapter is not a procedure and must never be able to act as one.
CHAPTER_KIND: str = "chapter"

#: The FDD specification. `manual` is one of the three kinds `DocumentChunk.kind` already
#: names, and it is the right one: manufacturer-or-engineer technical information, explicitly
#: distinguished from an SOP by `S4`, which refuses to answer a safety question from a manual.
SPECIFICATION_KIND: str = "manual"

#: The edition the reference document is on. **Sourced, not chosen** — `CONTEXT.md` §8 records
#: "Product & Architecture Document v4" and the file under `docs/00-source/` carries `_v4` in
#: its name. Every chapter split out of it inherits this and nothing else does.
REFERENCE_EDITION: str = "4"

#: A document whose source asserts no revision. Empty, and `DocumentChunk.cite` turns it into
#: the words `version not stated` — see `app/db/knowledge.VERSION_NOT_STATED`.
NO_VERSION_STATED: str = ""

#: How a chapter's document title is built. The collection first so a citation places the
#: passage before it names it, and the chapter's own H1 second so the address is the author's
#: rather than a filename.
PRODUCT_COLLECTION: str = "Synex Product & Architecture"

#: Architecture chapters numbered below this were written directly rather than split out of the
#: reference document, so they belong to no edition. `scripts/split_source.py` writes chapters
#: 26–44 into `docs/20-architecture/`; the four below that number are ours.
FIRST_SPLIT_ARCHITECTURE_CHAPTER: int = 26


@dataclass(frozen=True)
class SourceDocument:
    """One document, read verbatim, ready to be chunked.

    `digest` is taken over `text` at construction rather than by the caller, so a document
    cannot reach the store carrying a digest of something else — which would make the ingest
    silently skip a document that had changed.
    """

    title: str
    text: str
    kind: str
    origin: str
    """Where it was read from, repository-relative. What a reviewer opens to check a passage."""

    version: str = NO_VERSION_STATED

    @property
    def digest(self) -> str:
        return digest_of(self.text)

    @property
    def version_is_stated(self) -> bool:
        return bool(self.version.strip())

    def render(self) -> str:
        version = f"v{self.version}" if self.version_is_stated else "version not stated"
        return f"{self.title} ({version}, {self.kind}) — {self.origin}, {len(self.text)} chars"


@dataclass(frozen=True)
class RefusedSource:
    """A candidate this module decided **not** to ingest, and why in words.

    Public and reported on every run. A source silently not walked is indistinguishable from
    one nobody thought of, and the refusals here are the load-bearing half of §2.8 — each is a
    place the corpus would have grown a second source for a fact it already holds.
    """

    what: str
    reason: str

    def render(self) -> str:
        return f"{self.what} — not ingested: {self.reason}"


@dataclass(frozen=True)
class Discovery:
    """Everything the corpus is, and everything it deliberately is not."""

    documents: tuple[SourceDocument, ...] = field(default_factory=tuple)
    refused: tuple[RefusedSource, ...] = field(default_factory=tuple)
    unreadable: tuple[str, ...] = field(default_factory=tuple)
    """A file that was meant to be ingested and could not be read. **Never silently dropped**:
    a corpus missing a document it believes it holds scores well and answers wrongly."""

    @property
    def by_kind(self) -> dict[str, int]:
        counts: dict[str, int] = {}
        for document in self.documents:
            counts[document.kind] = counts.get(document.kind, 0) + 1
        return counts

    def render(self) -> str:
        lines = [f"{len(self.documents)} document(s) discovered: {self.by_kind}"]
        lines.extend(f"  {document.render()}" for document in self.documents)
        lines.extend(f"  {refusal.render()}" for refusal in self.refused)
        lines.extend(f"  UNREADABLE {problem}" for problem in self.unreadable)
        return "\n".join(lines)


# ── reading a file, and the one dependency that is not in requirements.txt ─────


def _title_from_markdown(text: str, fallback: str) -> str:
    """The document's own H1, or the filename if it has none.

    The heading rather than the filename because a citation is checked by a person against the
    document they are looking at, and what they see at the top of it is the heading.
    """
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("# "):
            return stripped[2:].strip()
        if stripped:
            break
    return fallback


def read_markdown(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def read_docx(path: Path) -> str:
    """Paragraphs and tables out of a `.docx`, in document order.

    **`python-docx` is imported here rather than at module scope, on purpose.** The offline
    test suite imports this module and must keep running with nothing installed beyond the
    gate's own dependencies; a module-level import would put a document parser in the path of
    every test that touches jobs. `scripts/split_source.py` already depends on the same
    library, so this adds a caller rather than a dependency.

    Headings are re-emitted as markdown hashes because that is the structure
    `app/retrieval/chunking.py` splits on. Nothing is reworded — a heading's own text is
    carried character for character and only the `#` marks are added, so a passage still reads
    as what the document says.
    """
    try:
        import docx  # noqa: PLC0415 — see the docstring
    except ImportError as exc:  # pragma: no cover — stated absence, never a silent skip
        raise RuntimeError(
            f"{path.name} is a .docx and python-docx is not installed, so it cannot be read. "
            f"That is a missing dependency rather than an empty document, and the two must not "
            f"look the same: pip install python-docx"
        ) from exc

    document = docx.Document(str(path))
    lines: list[str] = []
    for paragraph in document.paragraphs:
        content = paragraph.text.strip()
        if not content:
            continue
        style = (paragraph.style.name or "").lower()
        if style.startswith("heading"):
            level = "".join(c for c in style if c.isdigit()) or "1"
            lines.append(f"\n{'#' * min(int(level), 6)} {content}\n")
        elif style.startswith("title"):
            lines.append(f"\n# {content}\n")
        else:
            lines.append(content)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


# ── discovery ─────────────────────────────────────────────────────────────────


def _chapter_number(path: Path) -> int | None:
    """The numeric prefix `split_source.py` writes, or `None` if the file has none."""
    prefix = path.stem.split("-", 1)[0]
    return int(prefix) if prefix.isdigit() else None


def _chapter_version(directory: str, path: Path) -> str:
    """Edition 4 for a chapter split out of the reference document, nothing for one that was not.

    The rule is the numbering `scripts/split_source.py` produces: it writes the front matter,
    chapters 1–25 and Appendix A into `docs/10-product/` and chapters 26–44 into
    `docs/20-architecture/`. So everything in the product directory came out of edition 4, and
    in the architecture directory only the chapters at or above 26 did — the four below it are
    the data model, the stack, the deployment shape and the Thermynx inheritance, written
    directly and belonging to no edition.

    A file that cannot be placed by that rule gets **nothing** rather than the edition. Claiming
    a version a document never asserted is the failure `K5` names, and the cheap error here is
    a weaker citation.
    """
    if directory == "docs/10-product":
        return REFERENCE_EDITION
    number = _chapter_number(path)
    if number is not None and number >= FIRST_SPLIT_ARCHITECTURE_CHAPTER:
        return REFERENCE_EDITION
    return NO_VERSION_STATED


#: Directories of written chapters, and the file in each that is generated rather than written.
#: `README.md` in both is an index `scripts/split_source.py` regenerates — a table of contents
#: retrieved as a passage answers no question and cites a document it only lists.
_CHAPTER_DIRECTORIES: tuple[str, ...] = ("docs/10-product", "docs/20-architecture")
_GENERATED_INDEX: str = "README.md"


def discover(repo_root: Path) -> Discovery:
    """Walk the repository and decide, by name, what the corpus holds.

    Takes the root as an argument rather than computing it, so the whole judgement is testable
    against a directory a test builds — the property every gate in this repository holds.
    """
    documents: list[SourceDocument] = []
    unreadable: list[str] = []

    for directory in _CHAPTER_DIRECTORIES:
        folder = repo_root / directory
        if not folder.is_dir():
            unreadable.append(f"{directory} is not a directory")
            continue
        for path in sorted(folder.glob("*.md")):
            if path.name == _GENERATED_INDEX:
                continue
            try:
                text = read_markdown(path)
            except OSError as exc:
                unreadable.append(f"{directory}/{path.name}: {exc}")
                continue
            if not text.strip():
                unreadable.append(f"{directory}/{path.name}: the file is empty")
                continue
            documents.append(
                SourceDocument(
                    title=f"{PRODUCT_COLLECTION} — {_title_from_markdown(text, path.stem)}",
                    text=text,
                    kind=CHAPTER_KIND,
                    origin=f"{directory}/{path.name}",
                    version=_chapter_version(directory, path),
                )
            )

    specification = repo_root / "docs/00-source/WC_CHiller_FDD.docx"
    if specification.is_file():
        try:
            text = read_docx(specification)
        except (OSError, RuntimeError) as exc:
            unreadable.append(f"docs/00-source/{specification.name}: {exc}")
        else:
            documents.append(
                SourceDocument(
                    title="Water-cooled chiller FDD specification",
                    text=text,
                    kind=SPECIFICATION_KIND,
                    origin=f"docs/00-source/{specification.name}",
                    version=NO_VERSION_STATED,
                )
            )
    else:
        unreadable.append("docs/00-source/WC_CHiller_FDD.docx is absent")

    return Discovery(
        documents=tuple(documents), refused=refused_sources(), unreadable=tuple(unreadable)
    )


def refused_sources() -> tuple[RefusedSource, ...]:
    """Every candidate decided against, with the reason. Reported on every run.

    Constant rather than computed from the filesystem: the decision is about what a source
    *is*, not about whether it happens to be present, and a refusal that disappeared when
    somebody moved a file would be a decision nobody could review.
    """
    return (
        RefusedSource(
            what="the 78-page product and architecture reference (.docx)",
            reason=(
                "scripts/split_source.py splits it into the 47 chapter files that ARE "
                "ingested, so taking both puts every sentence in the corpus twice — "
                "CLAUDE.md §2.8. The .docx copy is also the worse one: it has no per-chapter "
                "address, so its citation would name a 78-page document and no place inside it"
            ),
        ),
        RefusedSource(
            what="the Feature Review Pack, in both .docx and .pdf",
            reason=(
                "its checklist content is transcribed into app/domain/library/ and is already "
                "the checklist corpus, so ingesting the pack would make three sources for one "
                "fact — and the .pdf and .docx are the same document twice over"
            ),
        ),
        RefusedSource(
            what="docs/90-archive/",
            reason=(
                "superseded editions. Retrieving one is the precise failure the version column "
                "exists to prevent: an instruction that was correct in an edition nobody "
                "follows any more"
            ),
        ),
        RefusedSource(
            what="mvp/, decisions/, brand/, CONTEXT.md, HANDOFF.md, CLAUDE.md",
            reason=(
                "documents about building the product rather than about the plant or the "
                "platform. A passage from the open-questions register retrieved into an answer "
                "would read as a statement about the equipment"
            ),
        ),
        RefusedSource(
            what="app/domain/library/holding_actions.py",
            reason=(
                "withheld by index_library.withheld_content() — nine drafted interim "
                "instructions behind two gates where DocumentChunk has one. Named here so the "
                "refusal is visible from the corpus as well as from the library ingest"
            ),
        ),
    )
